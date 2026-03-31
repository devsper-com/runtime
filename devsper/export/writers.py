from __future__ import annotations

from pathlib import Path
import html

from devsper.export.branding import Branding
from devsper.export.model import BundleExport, Citation, RunExport


def _safe_name(name: str) -> str:
    return "".join(ch if ch.isalnum() or ch in ("-", "_") else "_" for ch in name)[:80] or "run"


def _fmt_cost(v: float | None) -> str:
    return f"${v:.4f}" if v is not None else "-"


def _fmt_duration(seconds: float) -> str:
    s = max(0, int(seconds or 0))
    hh = s // 3600
    mm = (s % 3600) // 60
    ss = s % 60
    return f"{hh:02d}:{mm:02d}:{ss:02d}"


def _render_citations_md(citations: list[Citation]) -> str:
    if not citations:
        return "_No citations found in event history._\n"
    lines: list[str] = []
    for idx, c in enumerate(citations, 1):
        bits = [f"{idx}. "]
        bits.append(c.title or c.url or c.doi or c.arxiv_id or "reference")
        if c.doi:
            bits.append(f" DOI: {c.doi}")
        if c.arxiv_id:
            bits.append(f" arXiv:{c.arxiv_id}")
        if c.url:
            bits.append(f" ({c.url})")
        lines.append("".join(bits))
    return "\n".join(lines) + "\n"


def _render_single_run_md(run: RunExport) -> str:
    lines = [
        f"# Run {run.run_id}",
        "",
        f"- Root task: {run.root_task}",
        f"- Strategy: {run.strategy or '-'}",
        f"- Started: {run.started_at}",
        f"- Finished: {run.finished_at}",
        f"- Duration: {_fmt_duration(run.duration_seconds)}",
        f"- Tasks: {run.completed_tasks}/{run.total_tasks} completed, {run.failed_tasks} failed",
        f"- Estimated cost: {_fmt_cost(run.estimated_cost_usd)}",
        "",
        "## Tool usage",
    ]
    if run.tool_counts:
        for name, count in sorted(run.tool_counts.items(), key=lambda x: (-x[1], x[0])):
            lines.append(f"- {name}: {count}")
    else:
        lines.append("- none")

    lines += ["", "## Clarification history"]
    if run.clarifications:
        for qa in run.clarifications:
            lines.append(f"- [{qa.timestamp}] {qa.question}")
            lines.append(f"  - answer: {qa.answer}")
    else:
        lines.append("- none captured")

    lines += ["", "## Task outputs"]
    if run.task_outputs:
        for task_id, out in run.task_outputs.items():
            lines.append(f"### {task_id}")
            lines.append("")
            lines.append((out or "").strip() or "_empty_")
            lines.append("")
    else:
        lines.append("_Task outputs not present in event payloads for this run._")

    lines += ["", "## References / Citations", "", _render_citations_md(run.citations)]
    lines += ["## Timeline", ""]
    for item in run.timeline[:4000]:
        lines.append(f"- {item.timestamp} · {item.event_type} · {item.task_id} · {item.message}")
    return "\n".join(lines).strip() + "\n"


def render_bundle_markdown(bundle: BundleExport) -> str:
    lines = [
        "# devsper Run History Export",
        "",
        f"Generated at: {bundle.generated_at}",
        f"Total runs: {bundle.run_count}",
        "",
        "## Runs",
        "",
    ]
    for run in bundle.runs:
        lines.append(f"- `{run.run_id}` · {run.completed_tasks}/{run.total_tasks} completed · {_fmt_duration(run.duration_seconds)}")
    lines += ["", "## Global References", "", _render_citations_md(bundle.citations)]
    lines += ["## Run Details", ""]
    for run in bundle.runs:
        lines.append(f"### {run.run_id}")
        lines.append("")
        lines.append(f"- Root task: {run.root_task}")
        lines.append(f"- Duration: {_fmt_duration(run.duration_seconds)}")
        lines.append(f"- Tasks: {run.completed_tasks}/{run.total_tasks}, failed {run.failed_tasks}")
        lines.append(f"- Events path: {run.events_path or '(missing)'}")
        if run.tool_counts:
            top_tools = ", ".join(f"{k} x{v}" for k, v in sorted(run.tool_counts.items(), key=lambda x: -x[1])[:6])
            lines.append(f"- Top tools: {top_tools}")
        if run.clarifications:
            lines.append(f"- Clarifications: {len(run.clarifications)}")
        if run.task_outputs:
            first_tid = next(iter(run.task_outputs.keys()))
            preview = (run.task_outputs[first_tid] or "").strip().replace("\n", " ")
            if len(preview) > 220:
                preview = preview[:220] + "..."
            lines.append(f"- Output preview ({first_tid[:8]}): {preview}")
        lines.append(f"- Timeline events: {len(run.timeline)}")
        if run.task_outputs:
            lines.append("")
            lines.append("#### Task Outputs")
            lines.append("")
            for task_id, out in run.task_outputs.items():
                lines.append(f"- `{task_id}`")
                lines.append("")
                lines.append("```text")
                lines.append((out or "").strip()[:4000] or "_empty_")
                lines.append("```")
                lines.append("")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def render_bundle_rst(bundle: BundleExport) -> str:
    title = "devsper Run History Export"
    lines = [title, "=" * len(title), "", f"Generated at: {bundle.generated_at}", f"Total runs: {bundle.run_count}", ""]
    lines += ["Runs", "----", ""]
    for run in bundle.runs:
        lines.append(f"- ``{run.run_id}`` ({run.completed_tasks}/{run.total_tasks}, {_fmt_duration(run.duration_seconds)})")
    lines += ["", "Global References", "-----------------", ""]
    if bundle.citations:
        for idx, c in enumerate(bundle.citations, 1):
            ref = c.title or c.url or c.doi or c.arxiv_id or "reference"
            lines.append(f"{idx}. {ref}")
    else:
        lines.append("No citations found.")
    lines += ["", "Run Details", "-----------", ""]
    for run in bundle.runs:
        lines.append(f"{run.run_id}")
        lines.append("^" * len(run.run_id))
        lines.append(f"- Root task: {run.root_task}")
        lines.append(f"- Duration: {_fmt_duration(run.duration_seconds)}")
        lines.append(f"- Tasks: {run.completed_tasks}/{run.total_tasks}, failed {run.failed_tasks}")
        lines.append(f"- Events path: {run.events_path or '(missing)'}")
        lines.append(f"- Timeline events: {len(run.timeline)}")
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def _latex_escape(s: str) -> str:
    repl = {
        "\\": "\\textbackslash{}",
        "{": "\\{",
        "}": "\\}",
        "_": "\\_",
        "&": "\\&",
        "%": "\\%",
        "$": "\\$",
        "#": "\\#",
    }
    return "".join(repl.get(ch, ch) for ch in (s or ""))


def render_bundle_latex(bundle: BundleExport, branding: Branding) -> tuple[str, str]:
    tex_lines = [
        r"\documentclass{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{hyperref}",
        r"\title{" + _latex_escape(branding.product_name) + r" --- Run History Export}",
        r"\date{" + _latex_escape(bundle.generated_at) + "}",
        r"\begin{document}",
        r"\maketitle",
        r"\section*{Summary}",
        f"Total runs: {bundle.run_count}",
        r"\section*{Runs}",
        r"\begin{itemize}",
    ]
    for run in bundle.runs:
        tex_lines.append(
            r"\item "
            + _latex_escape(run.run_id)
            + f" ({run.completed_tasks}/{run.total_tasks}, {_fmt_duration(run.duration_seconds)})"
        )
    tex_lines += [r"\end{itemize}", r"\section*{Run Details}"]
    for run in bundle.runs:
        tex_lines.append(r"\subsection*{" + _latex_escape(run.run_id) + "}")
        tex_lines.append(r"\begin{itemize}")
        tex_lines.append(r"\item Root task: " + _latex_escape(run.root_task))
        tex_lines.append(r"\item Duration: " + _latex_escape(_fmt_duration(run.duration_seconds)))
        tex_lines.append(r"\item Tasks: " + _latex_escape(f"{run.completed_tasks}/{run.total_tasks}, failed {run.failed_tasks}"))
        tex_lines.append(r"\item Events path: " + _latex_escape(run.events_path or "(missing)"))
        tex_lines.append(r"\item Timeline events: " + _latex_escape(str(len(run.timeline))))
        tex_lines.append(r"\end{itemize}")
        if run.task_outputs:
            tex_lines.append(r"\paragraph{Task Outputs}")
            for task_id, out in run.task_outputs.items():
                tex_lines.append(r"\textbf{" + _latex_escape(task_id) + r"}\\")
                tex_lines.append(_latex_escape((out or "").strip()[:2000] or "_empty_"))
                tex_lines.append(r"\\")
    tex_lines += [r"\bibliographystyle{plain}", r"\bibliography{all_runs}", r"\end{document}"]

    bib_lines: list[str] = []
    for idx, c in enumerate(bundle.citations, 1):
        key = c.key or f"ref{idx}"
        title = _latex_escape(c.title or c.url or c.doi or c.arxiv_id or "reference")
        year = _latex_escape(c.year or "n.d.")
        author = _latex_escape(c.authors or "Unknown")
        bib_lines.append(f"@misc{{{key},")
        bib_lines.append(f"  author = {{{author}}},")
        bib_lines.append(f"  title = {{{title}}},")
        bib_lines.append(f"  year = {{{year}}},")
        if c.url:
            bib_lines.append(f"  url = {{{_latex_escape(c.url)}}},")
        elif c.doi:
            bib_lines.append(f"  note = {{DOI: {_latex_escape(c.doi)}}},")
        elif c.arxiv_id:
            bib_lines.append(f"  note = {{arXiv:{_latex_escape(c.arxiv_id)}}},")
        bib_lines.append("}\n")
    return ("\n".join(tex_lines) + "\n", "\n".join(bib_lines).strip() + "\n")


def render_bundle_html(bundle: BundleExport, branding: Branding) -> str:
    rows = []
    for run in bundle.runs:
        rows.append(
            f"<tr><td>{html.escape(run.run_id)}</td><td>{run.completed_tasks}/{run.total_tasks}</td>"
            f"<td>{_fmt_duration(run.duration_seconds)}</td><td>{html.escape(run.strategy or '-')}</td></tr>"
        )
    refs = []
    for c in bundle.citations:
        label = html.escape(c.title or c.url or c.doi or c.arxiv_id or "reference")
        refs.append(f"<li>{label}</li>")
    details = []
    for run in bundle.runs:
        output_blocks = []
        if run.task_outputs:
            for task_id, out in run.task_outputs.items():
                output_blocks.append(
                    f"<h4>Task {html.escape(task_id)}</h4><pre>{html.escape((out or '').strip()[:4000] or '_empty_')}</pre>"
                )
        outputs_html = "".join(output_blocks) if output_blocks else "<p><i>No task outputs captured.</i></p>"
        section = (
            "<section>"
            f"<h3>{html.escape(run.run_id)}</h3>"
            f"<p><b>Root task:</b> {html.escape(run.root_task)}</p>"
            f"<p><b>Duration:</b> {_fmt_duration(run.duration_seconds)} | <b>Tasks:</b> {run.completed_tasks}/{run.total_tasks} (failed {run.failed_tasks})</p>"
            f"<p><b>Events path:</b> {html.escape(run.events_path or '(missing)')}</p>"
            f"<p><b>Timeline events:</b> {len(run.timeline)}</p>"
            f"{outputs_html}"
            "</section>"
        )
        details.append(section)
    return f"""<!doctype html>
<html><head><meta charset="utf-8"><title>devsper export</title>
<style>
body {{ font-family: {branding.font_sans}, sans-serif; margin: 40px; color: #111; }}
h1 {{ color: {branding.primary_color}; }}
.meta {{ color: #555; }}
table {{ border-collapse: collapse; width: 100%; }}
td,th {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
</style></head><body>
<h1>{html.escape(branding.product_name)} - Run History Export</h1>
<p class="meta">Generated at {html.escape(bundle.generated_at)} | Runs {bundle.run_count}</p>
<h2>Runs</h2>
<table><thead><tr><th>Run ID</th><th>Tasks</th><th>Duration</th><th>Strategy</th></tr></thead>
<tbody>{''.join(rows)}</tbody></table>
<h2>References</h2><ol>{''.join(refs)}</ol>
<h2>Run Details</h2>{''.join(details)}
</body></html>"""


def write_bundle_files(bundle: BundleExport, out_dir: Path, branding: Branding) -> dict[str, str]:
    out_dir.mkdir(parents=True, exist_ok=True)
    runs_dir = out_dir / "runs"
    runs_dir.mkdir(parents=True, exist_ok=True)

    written: dict[str, str] = {}
    for run in bundle.runs:
        run_md = _render_single_run_md(run)
        p = runs_dir / f"{_safe_name(run.run_id)}.md"
        p.write_text(run_md, encoding="utf-8")
        written[f"run_md:{run.run_id}"] = str(p)

    bundle_md = out_dir / "all_runs.md"
    bundle_md.write_text(render_bundle_markdown(bundle), encoding="utf-8")
    written["all_runs_md"] = str(bundle_md)

    bundle_rst = out_dir / "all_runs.rst"
    bundle_rst.write_text(render_bundle_rst(bundle), encoding="utf-8")
    written["all_runs_rst"] = str(bundle_rst)

    tex, bib = render_bundle_latex(bundle, branding)
    tex_path = out_dir / "all_runs.tex"
    bib_path = out_dir / "all_runs.bib"
    tex_path.write_text(tex, encoding="utf-8")
    bib_path.write_text(bib, encoding="utf-8")
    written["all_runs_tex"] = str(tex_path)
    written["all_runs_bib"] = str(bib_path)

    html_path = out_dir / "all_runs.html"
    html_path.write_text(render_bundle_html(bundle, branding), encoding="utf-8")
    written["all_runs_html"] = str(html_path)

    # Optional DOCX (if python-docx is installed)
    try:
        from docx import Document  # type: ignore

        doc = Document()
        doc.add_heading(f"{branding.product_name} - Run History Export", level=0)
        doc.add_paragraph(f"Generated at {bundle.generated_at}")
        doc.add_paragraph(f"Total runs: {bundle.run_count}")
        doc.add_heading("Runs", level=1)
        for run in bundle.runs:
            doc.add_paragraph(
                f"{run.run_id} | tasks {run.completed_tasks}/{run.total_tasks} | {_fmt_duration(run.duration_seconds)} | {run.strategy or '-'}"
            )
        doc.add_heading("References", level=1)
        for c in bundle.citations:
            doc.add_paragraph(c.title or c.url or c.doi or c.arxiv_id or "reference")
        doc.add_heading("Run Details", level=1)
        for run in bundle.runs:
            doc.add_heading(run.run_id, level=2)
            doc.add_paragraph(f"Root task: {run.root_task}")
            doc.add_paragraph(f"Duration: {_fmt_duration(run.duration_seconds)}")
            doc.add_paragraph(f"Tasks: {run.completed_tasks}/{run.total_tasks}, failed {run.failed_tasks}")
            doc.add_paragraph(f"Events path: {run.events_path or '(missing)'}")
            doc.add_paragraph(f"Timeline events: {len(run.timeline)}")
            if run.task_outputs:
                doc.add_heading("Task Outputs", level=3)
                for task_id, out in run.task_outputs.items():
                    doc.add_paragraph(f"Task {task_id}")
                    doc.add_paragraph((out or "").strip()[:4000] or "_empty_")
        docx_path = out_dir / "all_runs.docx"
        doc.save(str(docx_path))
        written["all_runs_docx"] = str(docx_path)
    except Exception:
        pass

    return written
